import os
import re
from docx import Document
from docx.shared import Pt

def generate_report(data, po_nmber, po_folder_path ,template_path="resources\\templates\\template.docx"):

    # Opening the template file
    doc = Document(template_path)
    
    # Logging Surface test data
    # Decide if results passed pr failed and log them
    if data["Surface Test"]:
        results_surface = surface_test_results_check(data["Surface Test"])
        for i in range(len(results_surface["TSA"])):

            if results_surface["TSA"][i] == "Pass":
                modify_table_cell(doc, 0, i + 1, 2, "☑ Pass \xa0☐ Fail")
            else:
                modify_table_cell(doc, 0, i + 1, 2, "☐ Pass \xa0☑ Fail")

            if results_surface["SDA"][i] == "Pass":
                modify_table_cell(doc, 0, i + 4, 2, "☑ Pass \xa0☐ Fail")
            else:
                modify_table_cell(doc, 0, i + 4, 2, "☐ Pass \xa0☑ Fail")
        # Writre Test number
        modify_table_cell(doc, 0, 1, 1, data["Surface Test"]["Test Number"])

    # Logging Air test data
    # Decide if results passed pr failed and log them
    if data["Air Test"]:
        results_air = air_test_results_check(data["Air Test"])
        for i in range(len(results_air["TSA"])):

            if results_air["TSA"][i] == "Pass":
                modify_table_cell(doc, 0, i + 7, 2, "☑ Pass \xa0☐ Fail")
            else:
                modify_table_cell(doc, 0, i + 7, 2, "☐ Pass \xa0☑ Fail")

            if results_air["SDA"][i] == "Pass":
                modify_table_cell(doc, 0, i + 8, 2, "☑ Pass \xa0☐ Fail")
            else:
                modify_table_cell(doc, 0, i + 8, 2, "☐ Pass \xa0☑ Fail")
        # Writre Test number
        modify_table_cell(doc, 0, 7, 1, data["Air Test"]["Test Number"])

    # Logging Micro test data
    # Iterate over the samples
    if data["Microbiological Test"]:
        for i in range((len(data["Microbiological Test"]["Samples"]["LOT#"]))):
            # Write Lot and Report Numbers
            # Write the type of product and decide if the lot passed or not
            microbiological_results_check(doc, data["Microbiological Test"], i)

    # Logging Peel test data
    if data["Peel Test"]:
        for i in range(len(data["Peel Test"]["Results"])):
            peel_results_check(doc, data["Peel Test"], i)

    # Logging Cleaning batch data
    if data["Surface Test"] and data["Surface Test"]["Samples Data"]["Products"]:
        row_index = 0 # to keep track of the next row
        for i, product in enumerate(data["Surface Test"]["Samples Data"]["Products"]):
            for j, cat in enumerate(data["Surface Test"]["Samples Data"]["CAT#"][i]):
                for k, lot in enumerate(data["Surface Test"]["Samples Data"]["LOT#"][i][j]):
                    modify_table_cell(doc, 3, 2 + row_index, 1, product) # for product name
                    modify_table_cell(doc, 3, 2 + row_index, 0, cat) # for CAT#
                    modify_table_cell(doc, 3, 2 + row_index, 2, lot) # for LOT#
                    row_index += 1

    # report file path
    file_path = os.path.join(po_folder_path, f"PO {po_nmber}.docx")
    # Saving the report
    doc.save(file_path)
    print(f"Report generated and saved at {file_path}")

def modify_table_cell(doc : Document, table_index, row_index, col_index, new_text, font_size=12):
    """
    Modify the content of specific table cell in a docx file.

    :param doc:          Object of the .docx file
    :param table_index:  Integer index of the table you want to modify (0-based)
    :param row_index:    Integer index of the row you want to modify (0-based)
    :param col_index:    Integer index of the column you want to modify (0-based)
    :param new_text:     The new text you want to put in that cell
    """

    # Access the desired table (tables are 0-based)
    table = doc.tables[table_index]
    
    # Access the specific cell by row & column and set font to 10
    cell = table.cell(row_index, col_index)

    # Assign new text
    cell.text = str(new_text)
    selected_cell_run = cell.paragraphs[0].runs
    selected_cell_run[0].font.size = Pt(font_size)
    # paragraph = cell.paragraphs[0]
    # run = paragraph.add_run(str(new_text))
    # run.font.size = Pt(10)
    
def surface_test_results_check(data):

    # initializing test results dict
    test_results = {
        "TSA": [],
        "SDA": []
    }
    # Cnoverting results into integers, checking them and saving the pass/fail comdition in a way similar to the report
    for result in data["Results"]:
        for test in result["Tests"]:

            if "Total count- environmental sampling" in test["Test"]:
                environmental_sampling_match = re.findall(r"[^0-9]*([0-9]*)", test["Result"])[0]
                if environmental_sampling_match:
                    environmental_sampling_value = int(environmental_sampling_match)
                    if environmental_sampling_value > 50:
                        TSA = "Fail"
                    else :
                        TSA = "Pass"
            
            if "Yeasts - environmental tests" in test["Test"]:
                yeasts_match = re.findall(r"[^0-9]*([0-9]*)", test["Result"])[0]
                if yeasts_match:
                    yeasts_value = int(yeasts_match)
                    if yeasts_value > 50:
                        SDA_yeast = "Fail"
                    else:
                        SDA_yeast = "Pass"
            
            if "Molds" in test["Test"]:
                molds_match = re.findall(r"[^0-9]*([0-9]*)", test["Result"])[0]
                if molds_match:
                    molds_value = int(molds_match)
                    if molds_value > 50:
                        SDA_mold = "Fail"
                    else:
                        SDA_mold = "Pass"
            
        if TSA == "Pass":
            test_results["TSA"].append("Pass")
        else:
            test_results["TSA"].append("Fail")

        if SDA_mold == "Pass" and SDA_yeast == "Pass":
            test_results["SDA"].append("Pass")
        else:
            test_results["SDA"].append("Fail")

    return test_results

def air_test_results_check(data):

    # initializing test results dict
    test_results = {
        "TSA": [],
        "SDA": []
    }
    # Cnoverting results into integers, checking them and saving the pass/fail comdition in a way similar to the report
    for result in data["Results"]:
        for test in result["Tests"]:

            if "Total count- environmental sampling" in test["Test"]:
                environmental_sampling_match = re.findall(r"[^0-9]*([0-9]*)", test["Result"])[0]
                if environmental_sampling_match:
                    environmental_sampling_value = int(environmental_sampling_match)
                    if environmental_sampling_value > 100:
                        TSA = "Fail"
                    else :
                        TSA = "Pass"
            
            if "Yeasts - environmental tests" in test["Test"]:
                yeasts_match = re.findall(r"[^0-9]*([0-9]*)", test["Result"])[0]
                if yeasts_match:
                    yeasts_value = int(yeasts_match)
                    if yeasts_value > 100:
                        SDA_yeast = "Fail"
                    else:
                        SDA_yeast = "Pass"
            
            if "Molds" in test["Test"]:
                molds_match = re.findall(r"[^0-9]*([0-9]*)", test["Result"])[0]
                if molds_match:
                    molds_value = int(molds_match)
                    if molds_value > 100:
                        SDA_mold = "Fail"
                    else:
                        SDA_mold = "Pass"
            
        if TSA == "Pass":
            test_results["TSA"].append("Pass")
        else:
            test_results["TSA"].append("Fail")

        if SDA_mold == "Pass" and SDA_yeast == "Pass":
            test_results["SDA"].append("Pass")
        else:
            test_results["SDA"].append("Fail")

    return test_results

def microbiological_results_check(doc, data, index):

    # Write Test and Lot numbers
    modify_table_cell(doc, 1, 1 + index, 2, data["Samples"]["LOT#"][index])
    modify_table_cell(doc, 1, 1 + index, 3, data["Test Number"])

    # Saving results in one list
    results_list = []
    for i, result in enumerate(data["Results"]):
        result_match = re.findall(r"[^0-9]*([0-9]*)", result[1])
        results_list.append(int(result_match[0]))

    # Check product type, test result and writing to the table
    match data["Samples"]["Products"][0].lower():

        case "neuroprobe":
            product_string = "☑NeuroProbe \n☐Cannula \xa0\n☐ LeadConfirm cable\n☐ LeadConfirm Adaptor\n☐ AlphaProbe cable\n☐ Electrode cable"
            modify_table_cell(doc, 1, 1 + index, 1, product_string, font_size=10)
            if all(value <= 10 for value in results_list):
                modify_table_cell(doc, 1, 1 + index, 4, "☑ Pass \xa0☐ Fail")
            else:
                modify_table_cell(doc, 1, 1 + index, 4, "☐ Pass \xa0☑ Fail")

        case "cannula":
            product_string = "☐NeuroProbe \n☑Cannula \xa0\n☐ LeadConfirm cable\n☐ LeadConfirm Adaptor\n☐ AlphaProbe cable\n☐ Electrode cable"
            modify_table_cell(doc, 1, 1 + index, 1, product_string, font_size=10)
            if all(value <= 10 for value in results_list):
                modify_table_cell(doc, 1, 1 + index, 4, "☑ Pass \xa0☐ Fail")
            else:
                modify_table_cell(doc, 1, 1 + index, 4, "☐ Pass \xa0☑ Fail")

        case "leadconfirm cable":
            product_string = "☐NeuroProbe \n☐Cannula \xa0\n☑ LeadConfirm cable\n☐ LeadConfirm Adaptor\n☐ AlphaProbe cable\n☐ Electrode cable"
            modify_table_cell(doc, 1, 1 + index, 1, product_string, font_size=10)
            if all(value <= 100 for value in results_list):
                modify_table_cell(doc, 1, 1 + index, 4, "☑ Pass \xa0☐ Fail")
            else:
                modify_table_cell(doc, 1, 1 + index, 4, "☐ Pass \xa0☑ Fail")

        case "leadconfirm adaptor":
            product_string = "☐NeuroProbe \n☐Cannula \xa0\n☐ LeadConfirm cable\n☑ LeadConfirm Adaptor\n☐ AlphaProbe cable\n☐ Electrode cable"
            modify_table_cell(doc, 1, 1 + index, 1, product_string, font_size=10)
            if all(value <= 100 for value in results_list):
                modify_table_cell(doc, 1, 1 + index, 4, "☑ Pass \xa0☐ Fail")
            else:
                modify_table_cell(doc, 1, 1 + index, 4, "☐ Pass \xa0☑ Fail")

        case "alphaprobe cable":
            product_string = "☐NeuroProbe \n☐Cannula \xa0\n☐ LeadConfirm cable\n☐ LeadConfirm Adaptor\n☑ AlphaProbe cable\n☐ Electrode cable"
            modify_table_cell(doc, 1, 1 + index, 1, product_string, font_size=10)
            if all(value <= 100 for value in results_list):
                modify_table_cell(doc, 1, 1 + index, 4, "☑ Pass \xa0☐ Fail")
            else:
                modify_table_cell(doc, 1, 1 + index, 4, "☐ Pass \xa0☑ Fail")

        case "electrode cable":
            product_string = "☐NeuroProbe \n☐Cannula \xa0\n☐ LeadConfirm cable\n☐ LeadConfirm Adaptor\n☐ AlphaProbe cable\n☑ Electrode cable"
            modify_table_cell(doc, 1, 1 + index, 1, product_string, font_size=10)
            if all(value <= 100 for value in results_list):
                modify_table_cell(doc, 1, 1 + index, 4, "☑ Pass \xa0☐ Fail")
            else:
                modify_table_cell(doc, 1, 1 + index, 4, "☐ Pass \xa0☑ Fail")        

def peel_results_check(doc, data, index):

    # Write Test Number
    modify_table_cell(doc, 2, 1 + index, 1, data["Test Number"])

    # check if the value passes or fails
    if float(data["Results"][index][1]) > 1.5:
        modify_table_cell(doc, 2, 1 + index, 2, "☑ Pass ☐ Fail")
    else:
        modify_table_cell(doc, 2, 1 + index, 2, "☐ Pass ☑ Fail")
